import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import os
from datetime import datetime
from docx import Document

class LacrosseTimerApp:
    def __init__(self, root):
        self.root = root
        self.root.withdraw()  # hide the main window off the bat
        
        # show quarter length selection dialog first
        self.select_quarter_length()
        
        self.root.title("Lacrosse Timer App - © Dan Finn")
        self.root.configure(bg="#e6f2ff")  # for light blue background
        self.root.geometry("1200x800")
        
        # game clock variables
        self.quarter = 1
        # quarter_length is now set by select_quarter_length()
        self.game_clock_running = False
        self.game_clock_time = self.quarter_length
        self.game_paused_between_quarters = False
        
        # variables to track timers
        self.intervals = {}  # store timer objects
        self.paused_times = {}  # store paused times in seconds
        self.timer_frames = {}  # store timer frames
        self.timer_count = 2  # start with 2 timers by default
        self.timer_running_states = {}  # track which timers were running when quarter ended
        
        # create main frame with scrollbar
        self.main_frame = tk.Frame(root, bg="#e6f2ff")
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=70)
        
        # create game clock frame at the top
        self.create_game_clock()
        
        # create canvas with scrollbar
        self.canvas = tk.Canvas(self.main_frame, bg="#e6f2ff")
        self.scrollbar = ttk.Scrollbar(self.main_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # create a frame inside the canvas for the timers
        self.timer_container = tk.Frame(self.canvas, bg="#e6f2ff")
        self.canvas.create_window((0, 0), window=self.timer_container, anchor="nw")
        
        # header
        self.header = tk.Frame(root, bg="#007BFF", height=50)
        self.header.pack(fill=tk.X, side=tk.TOP)
        self.header.pack_propagate(False)
        
        self.title_label = tk.Label(
            self.header, 
            text="Lacrosse Timer App - © Dan Finn",
            font=("Arial", 18),
            bg="#007BFF",
            fg="white"
        )
        self.title_label.pack(pady=10, side=tk.LEFT, padx=10)
        
        # exit button in top right
        self.exit_btn = tk.Button(
            self.header,
            text="✕ Exit",
            command=self.exit_application,
            font=("Arial", 12, "bold"),
            bg="#FF5252",
            fg="white",
            bd=0,
            padx=15,
            pady=5
        )
        self.exit_btn.pack(side=tk.RIGHT, padx=10)
        
        # settings button
        self.settings_btn = tk.Button(
            self.header,
            text="⚙️ Settings",
            command=self.open_settings,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0,
            padx=10,
            pady=5
        )
        self.settings_btn.pack(side=tk.RIGHT, padx=10)
        
        # footer with control buttons
        self.footer = tk.Frame(root, bg="#007BFF", height=50)
        self.footer.pack(fill=tk.X, side=tk.BOTTOM)
        self.footer.pack_propagate(False)
        
        self.start_all_btn = tk.Button(
            self.footer, 
            text="Start All", 
            command=self.start_all_timers,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.start_all_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        self.stop_all_btn = tk.Button(
            self.footer, 
            text="Stop All", 
            command=self.stop_all_timers,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.stop_all_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        self.resume_all_btn = tk.Button(
            self.footer, 
            text="Resume All", 
            command=self.resume_all_timers,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.resume_all_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        # add & remove timer buttons
        self.add_timer_btn = tk.Button(
            self.footer, 
            text="Add Timer", 
            command=self.add_timer,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.add_timer_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        self.remove_timer_btn = tk.Button(
            self.footer, 
            text="Remove Timer", 
            command=self.remove_timer,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.remove_timer_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        # save button
        self.save_btn = tk.Button(
            self.footer, 
            text="Save", 
            command=self.save_data,
            font=("Arial", 12),
            bg="#007BFF",
            fg="white",
            bd=0
        )
        self.save_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5, pady=10)
        
        # initialize timers
        self.initialize_timers()
        
        # load saved data if exists
        self.load_data()
        
        # configure canvas scrolling
        self.timer_container.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)
        
        # bind mousewheel for scrolling
        self.canvas.bind_all("<MouseWheel>", self.on_mousewheel)
        
        # bind fullscreen toggle
        self.root.bind("<F11>", self.toggle_fullscreen)
        self.root.bind("<Escape>", self.end_fullscreen)
        
        # set up window size constraints
        self.root.update()
        self.root.minsize(800, 600)
        
        # adjust timer window sizes for fullscreen
        self.root.bind("<Configure>", self.adjust_timer_sizes)
        
        # show the main window now that everything is set up
        self.root.deiconify()
    
    def select_quarter_length(self):
        """Show dialog to select quarter length"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Select Quarter Length")
        dialog.geometry("300x200")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # center the dialog
        dialog.update_idletasks()
        width = dialog.winfo_width()
        height = dialog.winfo_height()
        x = (dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (dialog.winfo_screenheight() // 2) - (height // 2)
        dialog.geometry(f'{width}x{height}+{x}+{y}')
        
        tk.Label(
            dialog,
            text="Select Quarter Length:",
            font=("Arial", 14, "bold")
        ).pack(pady=(20, 10))
        
        # create time options from 2 to 20 minutes in 2-minute intervals
        times = [f"{i} minutes" for i in range(2, 21, 2)]
        time_var = tk.StringVar(value="12 minutes")  # default to 12 minutes
        
        time_dropdown = ttk.Combobox(
            dialog,
            textvariable=time_var,
            values=times,
            state="readonly",
            font=("Arial", 12),
            width=15
        )
        time_dropdown.pack(pady=20)
        
        def set_time():
            selected = time_var.get()
            minutes = int(selected.split()[0])
            self.quarter_length = minutes * 60
            dialog.destroy()
        
        tk.Button(
            dialog,
            text="Start Game",
            command=set_time,
            font=("Arial", 12, "bold"),
            bg="#007BFF",
            fg="white",
            padx=20,
            pady=10
        ).pack(pady=20)
        
        # wait for the dialog to be closed
        self.root.wait_window(dialog)
    
    def seconds_to_ms(self, seconds):
        """Convert seconds to MM:SS format"""
        minutes = seconds // 60
        secs = seconds % 60
        return f"{minutes:02d}:{secs:02d}"
    
    def ms_to_seconds(self, ms_str):
        """Convert MM:SS string to seconds"""
        try:
            minutes, seconds = map(int, ms_str.split(":"))
            return minutes * 60 + seconds
        except:
            return 0
    
    def seconds_to_hms(self, seconds):
        """Convert seconds to HH:MM:SS format"""
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        secs = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def hms_to_seconds(self, hms_str):
        """Convert HH:MM:SS string to seconds"""
        try:
            parts = hms_str.split(":")
            if len(parts) == 3:
                hours, minutes, seconds = map(int, parts)
                return hours * 3600 + minutes * 60 + seconds
            elif len(parts) == 2:
                minutes, seconds = map(int, parts)
                return minutes * 60 + seconds
            else:
                return 0
        except:
            return 0
    
    def create_game_clock(self):
        """Create the game clock display at the top of the app"""
        self.game_clock_frame = tk.Frame(self.root, bg="#004080", bd=2, relief=tk.RAISED)
        self.game_clock_frame.place(relx=0.5, y=60, anchor="n", width=400, height=120)
        
        # quarter display
        self.quarter_frame = tk.Frame(self.game_clock_frame, bg="#004080")
        self.quarter_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        self.quarter_label = tk.Label(
            self.quarter_frame,
            text=f"Quarter: {self.quarter}/4",
            font=("Arial", 16, "bold"),
            bg="#004080",
            fg="white"
        )
        self.quarter_label.pack(side=tk.LEFT, padx=10)
        
        # game clock display
        self.game_clock_display = tk.Label(
            self.game_clock_frame,
            text=self.seconds_to_ms(self.game_clock_time),
            font=("Arial", 36, "bold"),
            bg="#004080",
            fg="white"
        )
        self.game_clock_display.pack(pady=5)
        
        # game clock controls
        self.game_clock_controls = tk.Frame(self.game_clock_frame, bg="#004080")
        self.game_clock_controls.pack(side=tk.BOTTOM, fill=tk.X, pady=5)
        
        self.start_game_btn = tk.Button(
            self.game_clock_controls,
            text="Start",
            command=self.start_game_clock,
            bg="#00CC66",
            fg="white",
            font=("Arial", 12, "bold"),
            width=6
        )
        self.start_game_btn.pack(side=tk.LEFT, padx=5)
        
        self.stop_game_btn = tk.Button(
            self.game_clock_controls,
            text="Stop",
            command=self.stop_game_clock,
            bg="#FF5252",
            fg="white",
            font=("Arial", 12, "bold"),
            width=6
        )
        self.stop_game_btn.pack(side=tk.LEFT, padx=5)
        
        self.next_quarter_btn = tk.Button(
            self.game_clock_controls,
            text="Next Quarter",
            command=self.next_quarter,
            bg="#FFA500",
            fg="white",
            font=("Arial", 12, "bold")
        )
        self.next_quarter_btn.pack(side=tk.LEFT, padx=5)
    
    def start_game_clock(self):
        """Start the game clock"""
        if not self.game_clock_running:
            self.game_clock_running = True
            self.update_game_clock()
            
            # if we were paused between quarters, resume penalty timers
            if self.game_paused_between_quarters:
                self.game_paused_between_quarters = False
                for index, was_running in self.timer_running_states.items():
                    if was_running and index in self.timer_frames:
                        self.start_timer(index)
    
    def stop_game_clock(self):
        """Stop the game clock"""
        self.game_clock_running = False
        if hasattr(self, 'game_clock_after_id'):
            self.root.after_cancel(self.game_clock_after_id)
    
    def update_game_clock(self):
        """Update the game clock display"""
        if self.game_clock_running and self.game_clock_time > 0:
            self.game_clock_time -= 1
            self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
            self.game_clock_after_id = self.root.after(1000, self.update_game_clock)
        elif self.game_clock_running and self.game_clock_time <= 0:
            # quarter has ended
            self.game_clock_running = False
            if self.quarter < 4:
                messagebox.showinfo("Quarter End", f"Quarter {self.quarter} has ended!")
                # store which timers were running
                self.timer_running_states = {}
                for index in list(self.intervals.keys()):
                    self.timer_running_states[index] = True
                    self.stop_timer(index)
                self.game_paused_between_quarters = True
            else:
                # game over
                self.handle_game_over()
    
    def next_quarter(self):
        """Move to the next quarter"""
        if self.quarter < 4:
            self.quarter += 1
            self.game_clock_time = self.quarter_length
            self.quarter_label.config(text=f"Quarter: {self.quarter}/4")
            self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
            messagebox.showinfo("Next Quarter", f"Starting Quarter {self.quarter}")
        else:
            messagebox.showinfo("Game Over", "The game is already in the final quarter!")
    
    def handle_game_over(self):
        """Handle end of game actions"""
        response = messagebox.askyesno("Game Over", "Game over, do you want to save the game data to a Word file?")
        if response:
            self.export_to_word()
    
    def add_timer(self):
        """Add a new timer"""
        new_index = max(self.timer_frames.keys(), default=0) + 1
        self.create_timer(new_index)
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def remove_timer(self):
        """Remove the last timer"""
        if len(self.timer_frames) <= 1:
            messagebox.showinfo("Cannot Remove", "You must have at least one timer.")
            return
        
        last_index = max(self.timer_frames.keys())
        self.remove_specific_timer(last_index)

    def stop_timer(self, index):
        """Stop the timer with the given index"""
        if index in self.intervals:
            self.root.after_cancel(self.intervals[index])
            del self.intervals[index]

    def start_all_timers(self):
        """Start all active timers"""
        if not self.game_clock_running:
            messagebox.showinfo("Game Clock Stopped", "Please start the game clock first.")
            return

        for index in self.timer_frames:
            if self.paused_times[index] > 0:
                self.start_timer(index)

    def stop_all_timers(self):
        """Stop all running timers"""
        for index in list(self.intervals.keys()):
            self.stop_timer(index)

    def resume_all_timers(self):
        """Resume all paused timers"""
        if not self.game_clock_running:
            messagebox.showinfo("Game Clock Stopped", "Please start the game clock first.")
            return
        
        for index in self.timer_frames:
            if self.paused_times[index] > 0:
                self.start_timer(index)

    def save_data(self):
        """Save the app state to a JSON file"""
        try:
            data = {
                "quarter": self.quarter,
                "quarter_length": self.quarter_length,
                "game_clock_time": self.game_clock_time,
                "timers": {}
            }

            # save timer data
            for index, timer_data in self.timer_frames.items():
                data["timers"][str(index)] = {
                    "player_number": timer_data["player_entry"].get(),
                    "team_name": timer_data["team_entry"].get(),
                    "penalty_type": timer_data["penalty_var"].get(),
                    "penalty_time": timer_data["penalty_time_entry"].get(),
                    "time_option": timer_data["time_var"].get(),
                    "paused_time": self.paused_times[index]
                }

            with open("lacrosse_timer_data.json", "w") as f:
                json.dump(data, f)

            messagebox.showinfo("Save Successful", "Game data has been saved.")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save data: {str(e)}")

    def export_to_word(self):
        """Export all game data to a Word document"""
        try:
            # ask user where to save the file
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                title="Save Game Data"
            )
            
            if not file_path:
                return  # sser cancelled
            
            # create a new Word document
            doc = Document()
            
            # add title
            doc.add_heading('Lacrosse Game Report', 0)
            
            # add date and time
            doc.add_paragraph(f"Game Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # add a horizontal line
            doc.add_paragraph('_' * 50)
            
            # game summary section
            doc.add_heading('Game Summary', level=1)
            doc.add_paragraph(f"Total Quarters Played: {self.quarter}")
            doc.add_paragraph(f"Quarter Length: {self.quarter_length // 60} minutes")
            
            # penalty summary
            doc.add_heading('Penalty Summary', level=1)
            
            # create a table for penalties
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Player Number'
            header_cells[1].text = 'Team Name'
            header_cells[2].text = 'Penalty Type'
            header_cells[3].text = 'Penalty Time'
            header_cells[4].text = 'Penalty Duration'
            
            # add data rows
            for index, timer_data in self.timer_frames.items():
                player_number = timer_data["player_entry"].get()
                team_name = timer_data["team_entry"].get()
                penalty_type = timer_data["penalty_var"].get()
                penalty_time = timer_data["penalty_time_entry"].get()
                penalty_duration = timer_data["time_var"].get()
                
                # only add rows with actual data
                if player_number or team_name or penalty_type != "Select Penalty Type":
                    row_cells = table.add_row().cells
                    row_cells[0].text = player_number
                    row_cells[1].text = team_name
                    row_cells[2].text = penalty_type
                    row_cells[3].text = penalty_time
                    row_cells[4].text = penalty_duration
            
            # add footer
            doc.add_paragraph('_' * 50)
            doc.add_paragraph('Generated by Lacrosse Timer App - © Dan Finn')
            
            # save the document
            doc.save(file_path)
            messagebox.showinfo("Export Successful", f"Game data has been exported to {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export data: {str(e)}")
    
    def open_settings(self):
        """Open the settings dialog"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Settings")
        settings_window.geometry("400x450")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        # quarter length setting
        quarter_frame = tk.Frame(settings_window, pady=10)
        quarter_frame.pack(fill=tk.X, padx=20)
        
        tk.Label(quarter_frame, text="Quarter Length (minutes):").pack(side=tk.LEFT)
        
        quarter_length_var = tk.StringVar(value=str(self.quarter_length // 60))
        quarter_entry = tk.Entry(quarter_frame, textvariable=quarter_length_var, width=5)
        quarter_entry.pack(side=tk.LEFT, padx=10)
        
        # time adjustment controls
        adjust_frame = tk.LabelFrame(settings_window, text="Manual Time Override", pady=10, padx=10)
        adjust_frame.pack(fill=tk.X, padx=20, pady=10)
        
        tk.Label(adjust_frame, text="Adjust all timers by:").grid(row=0, column=0, sticky="w", pady=5)
        
        adjust_var = tk.StringVar(value="0")
        adjust_entry = tk.Entry(adjust_frame, textvariable=adjust_var, width=5)
        adjust_entry.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(adjust_frame, text="seconds").grid(row=0, column=2, sticky="w", pady=5)
        
        btn_frame = tk.Frame(adjust_frame)
        btn_frame.grid(row=1, column=0, columnspan=3, pady=10)
        
        tk.Button(
            btn_frame, 
            text="Add Time", 
            command=lambda: self.adjust_all_timers(int(adjust_var.get())),
            bg="#00CC66",
            fg="white"
        ).pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            btn_frame, 
            text="Subtract Time", 
            command=lambda: self.adjust_all_timers(-int(adjust_var.get())),
            bg="#FF5252",
            fg="white"
        ).pack(side=tk.LEFT, padx=5)
        
        # game management section
        game_frame = tk.LabelFrame(settings_window, text="Game Management", pady=10, padx=10)
        game_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # start New Game button
        tk.Button(
            game_frame,
            text="Start New Game",
            command=self.start_new_game,
            bg="#007BFF",
            fg="white",
            font=("Arial", 12)
        ).pack(fill=tk.X, pady=5)
        
        # export to Word button
        tk.Button(
            game_frame,
            text="Export All Data to Word",
            command=self.export_to_word,
            bg="#007BFF",
            fg="white",
            font=("Arial", 12)
        ).pack(fill=tk.X, pady=5)
        
        # clear Memory button
        tk.Button(
            game_frame,
            text="Clear Current Page Memory",
            command=self.clear_memory,
            bg="#FF5252",
            fg="white",
            font=("Arial", 12)
        ).pack(fill=tk.X, pady=5)
        
        # save settings button
        save_frame = tk.Frame(settings_window, pady=10)
        save_frame.pack(fill=tk.X, padx=20, pady=10)
        
        tk.Button(
            save_frame,
            text="Save Settings",
            command=lambda: self.save_settings(int(quarter_length_var.get()), settings_window),
            bg="#007BFF",
            fg="white",
            font=("Arial", 12, "bold")
        ).pack(fill=tk.X)
    
    def save_settings(self, quarter_length_minutes, settings_window):
        """Save the settings and close the settings window"""
        self.quarter_length = quarter_length_minutes * 60
        
        # if we're in a new quarter, update the time
        if self.game_clock_time == self.quarter_length or self.game_clock_time == 0:
            self.game_clock_time = self.quarter_length
            self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
        
        settings_window.destroy()
        messagebox.showinfo("Settings Saved", "Your settings have been saved.")
    
    def adjust_all_timers(self, seconds_to_adjust):
        """Adjust all timers by the specified number of seconds"""
        # adjust game clock
        self.game_clock_time = max(0, self.game_clock_time + seconds_to_adjust)
        self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
        
        # adjust penalty timers
        for index in self.paused_times:
            if self.paused_times[index] > 0:  # only adjust active timers
                self.paused_times[index] = max(0, self.paused_times[index] + seconds_to_adjust)
                if index in self.timer_frames:
                    self.timer_frames[index]["time_display"].config(text=self.seconds_to_hms(self.paused_times[index]))
        
        # save the updated state
        self.save_data()
    
    def start_new_game(self):
        """Start a new game with fresh settings"""
        confirm = messagebox.askyesno("Confirm New Game", "Are you sure you want to start a new game? This will reset all timers and game data.")
        if confirm:
            # stop all timers
            self.stop_game_clock()
            self.stop_all_timers()
            
            # show quarter length selection dialog
            self.select_quarter_length()
            
            # reset game state
            self.quarter = 1
            self.game_clock_time = self.quarter_length
            self.game_paused_between_quarters = False
            
            # update game clock display
            self.quarter_label.config(text=f"Quarter: {self.quarter}/4")
            self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
            
            # clear all timers
            for index in list(self.timer_frames.keys()):
                self.timer_frames[index]["frame"].destroy()
            
            self.timer_frames = {}
            self.intervals = {}
            self.paused_times = {}
            self.timer_running_states = {}
            
            # initialize new timers
            self.initialize_timers()
            
            # save the new state
            self.save_data()
            
            messagebox.showinfo("New Game", "New game has been started.")
    
    def clear_memory(self):
        """Clear all current page data without starting a new game"""
        confirm = messagebox.askyesno("Confirm Clear Memory", "Are you sure you want to clear all current timer data? This will not reset the game clock or quarter.")
        if confirm:
            # stop all timers
            self.stop_all_timers()
            
            # clear all timers
            for index in list(self.timer_frames.keys()):
                self.timer_frames[index]["frame"].destroy()
            
            self.timer_frames = {}
            self.intervals = {}
            self.paused_times = {}
            self.timer_running_states = {}
            
            # initialize new timers
            self.initialize_timers()
            
            # save the new state
            self.save_data()
            
            messagebox.showinfo("Memory Cleared", "All timer data has been cleared.")
    
    def exit_application(self):
        """Save data and exit the application"""
        confirm = messagebox.askyesno("Confirm Exit", "Are you sure you want to exit? Your data will be saved automatically.")
        if confirm:
            self.save_data()
            self.root.destroy()
    
    def toggle_fullscreen(self, event=None):
        """Toggle fullscreen mode"""
        state = not self.root.attributes("-fullscreen")
        self.root.attributes("-fullscreen", state)
        if state:
            # adjust timer sizes for fullscreen
            for timer_frame in self.timer_frames.values():
                timer_frame["frame"].config(width=400)  # set a fixed width for timers in fullscreen
        return "break"
    
    def end_fullscreen(self, event=None):
        """End fullscreen mode"""
        self.root.attributes("-fullscreen", False)
        return "break"
    
    def adjust_timer_sizes(self, event=None):
        """Adjust timer sizes based on window size"""
        if hasattr(self, 'timer_frames'):
            window_width = self.root.winfo_width()
            
            # if window is very wide, limit timer width
            if window_width > 1200:
                max_width = min(600, window_width // 2)
                for index in self.timer_frames:
                    self.timer_frames[index]["frame"].config(width=max_width)
    
    def setup_timer(self, index):
        """Set up the timer with the selected time"""
        if index in self.timer_frames:
            time_str = self.timer_frames[index]["time_var"].get()
            if time_str != "Not in use":
                seconds = self.hms_to_seconds(time_str)
                self.paused_times[index] = seconds
                self.timer_frames[index]["time_display"].config(text=time_str)
            
                # auto-fill penalty time with current game clock time
                if not self.timer_frames[index]["penalty_time_entry"].get():
                    game_time = self.seconds_to_ms(self.game_clock_time)
                    self.timer_frames[index]["penalty_time_entry"].delete(0, tk.END)
                    self.timer_frames[index]["penalty_time_entry"].insert(0, game_time)
            else:
                self.paused_times[index] = 0
                self.timer_frames[index]["time_display"].config(text="00:00:00")

    def released_timer(self, index):
        """Mark a player as released and clear the timer"""
        if index in self.timer_frames:
            # stop the timer if running
            self.stop_timer(index)
        
            # reset the timer display
            self.paused_times[index] = 0
            self.timer_frames[index]["time_display"].config(text="00:00:00")
            self.timer_frames[index]["time_var"].set("Not in use")
        
            # add "Released" to the player number
            current_player = self.timer_frames[index]["player_entry"].get()
            if current_player and "Released" not in current_player:
                self.timer_frames[index]["player_entry"].delete(0, tk.END)
                self.timer_frames[index]["player_entry"].insert(0, f"{current_player} (Released)")
            
            # save the updated state
            self.save_data()

    def remove_specific_timer(self, index):
        """Remove a specific timer by its index"""
        if len(self.timer_frames) <= 1:
            messagebox.showinfo("Cannot Remove", "You must have at least one timer.")
            return
        
        # stop the timer if running
        self.stop_timer(index)
    
        # remove the timer frame
        if index in self.timer_frames:
            self.timer_frames[index]["frame"].destroy()
            del self.timer_frames[index]
        
            # remove from paused_times
            if index in self.paused_times:
                del self.paused_times[index]
        
            # update scroll region
            self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            
            # save the updated state
            self.save_data()

    def start_timer(self, index):
        """Start the timer with the given index"""
        if index not in self.timer_frames or not self.game_clock_running:
            return
    
        # stop existing timer if running
        self.stop_timer(index)
    
        # start a new timer
        if self.paused_times[index] > 0:
            self.intervals[index] = self.root.after(1000, lambda: self.update_timer(index))

    def update_timer(self, index):
        """Update the timer display and decrement the time"""
        if index not in self.timer_frames or index not in self.paused_times:
            return
    
        # only update if game clock is running
        if self.game_clock_running and self.paused_times[index] > 0:
            self.paused_times[index] -= 1
            self.timer_frames[index]["time_display"].config(text=self.seconds_to_hms(self.paused_times[index]))
        
            # check if timer has reached zero
            if self.paused_times[index] <= 0:
                # timer completed
                self.stop_timer(index)
                # play a sound or flash the timer to indicate completion
                self.timer_frames[index]["frame"].config(bg="#FFCCCC")  # light red background
                self.root.after(3000, lambda: self.timer_frames[index]["frame"].config(bg="white"))  # reset after 3 seconds
            else:
                # continue the timer
                self.intervals[index] = self.root.after(1000, lambda: self.update_timer(index))
        elif not self.game_clock_running:
            # pause timer if game clock is stopped
            self.stop_timer(index)

    def load_data(self):
        """Load the app state from a JSON file"""
        try:
            if not os.path.exists("lacrosse_timer_data.json"):
                return
        
            with open("lacrosse_timer_data.json", "r") as f:
                data = json.load(f)
        
            # load game state
            self.quarter = data.get("quarter", 1)
            self.quarter_length = data.get("quarter_length", 12 * 60)
            self.game_clock_time = data.get("game_clock_time", self.quarter_length)
        
            # update game clock display
            self.quarter_label.config(text=f"Quarter: {self.quarter}/4")
            self.game_clock_display.config(text=self.seconds_to_ms(self.game_clock_time))
        
            # clear existing timers
            for index in list(self.timer_frames.keys()):
                self.timer_frames[index]["frame"].destroy()
            self.timer_frames.clear()
            self.paused_times.clear()
        
            # load timer data
            for index_str, timer_data in data.get("timers", {}).items():
                index = int(index_str)
                self.create_timer(index)
            
                # set timer values
                self.timer_frames[index]["player_entry"].insert(0, timer_data.get("player_number", ""))
                self.timer_frames[index]["team_entry"].insert(0, timer_data.get("team_name", ""))
                self.timer_frames[index]["penalty_var"].set(timer_data.get("penalty_type", "Select Penalty Type"))
                self.timer_frames[index]["penalty_time_entry"].insert(0, timer_data.get("penalty_time", ""))
                self.timer_frames[index]["time_var"].set(timer_data.get("time_option", "Not in use"))
            
                # set paused time
                self.paused_times[index] = timer_data.get("paused_time", 0)
                self.timer_frames[index]["time_display"].config(text=self.seconds_to_hms(self.paused_times[index]))
        
            messagebox.showinfo("Load Successful", "Game data has been loaded.")
        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load data: {str(e)}")

    def on_frame_configure(self, event):
        """Reset the scroll region to encompass the inner frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def on_canvas_configure(self, event):
        """When canvas is resized, resize the inner frame to match"""
        width = event.width
        self.canvas.itemconfig(self.canvas.find_withtag("all")[0], width=width)
    
    def on_mousewheel(self, event):
        """Handle mousewheel scrolling"""
        self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    
    def initialize_timers(self):
        """Create initial timers"""
        for i in range(1, self.timer_count + 1):
            self.create_timer(i)
    
    def create_timer(self, index):
        """Create a new timer with the given index"""
        # create a frame for the timer
        timer_frame = tk.Frame(self.timer_container, bg="white", bd=1, relief=tk.SOLID, padx=10, pady=10)
        timer_frame.pack(fill=tk.X, padx=10, pady=10)
        
        # add close button (X) to remove this specific timer
        close_btn = tk.Button(
            timer_frame,
            text="✕",
            command=lambda idx=index: self.remove_specific_timer(idx),
            bg="white",
            fg="red",
            font=("Arial", 10, "bold"),
            bd=0,
            padx=5,
            pady=0
        )
        close_btn.grid(row=0, column=2, sticky="ne")
        
        # player number input
        player_label = tk.Label(timer_frame, text="Player Number:", bg="white")
        player_label.grid(row=0, column=0, sticky="w", padx=5, pady=5)
        player_entry = tk.Entry(timer_frame, width=20)
        player_entry.grid(row=0, column=1, sticky="w", padx=5, pady=5)
        
        # team name input
        team_label = tk.Label(timer_frame, text="Team Name:", bg="white")
        team_label.grid(row=1, column=0, sticky="w", padx=5, pady=5)
        team_entry = tk.Entry(timer_frame, width=20)
        team_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)
        
        # time selection
        time_label = tk.Label(timer_frame, text="Penalty Time:", bg="white")
        time_label.grid(row=2, column=0, sticky="w", padx=5, pady=5)
        
        time_options = [
            "Not in use",
            "00:00:30",
            "00:01:00",
            "00:01:30",
            "00:02:00",
            "00:02:30",
            "00:03:00",
            "00:03:30",
            "00:04:00",
            "00:04:30",
            "00:05:00"
        ]
        
        time_var = tk.StringVar(value=time_options[0])
        time_dropdown = ttk.Combobox(timer_frame, textvariable=time_var, values=time_options, state="readonly", width=18)
        time_dropdown.grid(row=2, column=1, sticky="w", padx=5, pady=5)
        time_dropdown.bind("<<ComboboxSelected>>", lambda e, idx=index: self.setup_timer(idx))
        
        # timer display
        time_display = tk.Label(timer_frame, text="00:00:00", font=("Arial", 16, "bold"), bg="white")
        time_display.grid(row=3, column=0, columnspan=2, pady=10)
        
        # control buttons
        button_frame = tk.Frame(timer_frame, bg="white")
        button_frame.grid(row=4, column=0, columnspan=2, pady=5)
        
        start_btn = tk.Button(
            button_frame, 
            text="Start", 
            command=lambda idx=index: self.start_timer(idx),
            bg="#007BFF",
            fg="white"
        )
        start_btn.pack(side=tk.LEFT, padx=5)
        
        stop_btn = tk.Button(
            button_frame, 
            text="Stop", 
            command=lambda idx=index: self.stop_timer(idx),
            bg="#007BFF",
            fg="white"
        )
        stop_btn.pack(side=tk.LEFT, padx=5)
        
        released_btn = tk.Button(
            button_frame, 
            text="Released", 
            command=lambda idx=index: self.released_timer(idx),
            bg="#007BFF",
            fg="white"
        )
        released_btn.pack(side=tk.LEFT, padx=5)
        
        # penalty type selection
        penalty_label = tk.Label(timer_frame, text="Penalty Type:", bg="white")
        penalty_label.grid(row=5, column=0, sticky="w", padx=5, pady=5)
        
        penalty_options = [
            "Select Penalty Type",
            "Foul",
            "Slash",
            "Push",
            "Crease",
            "Crosscheck",
            "IllegalProcedure",
            "Screen",
            "Interference",
            "Misconduct"
        ]
        
        penalty_var = tk.StringVar(value=penalty_options[0])
        penalty_dropdown = ttk.Combobox(timer_frame, textvariable=penalty_var, values=penalty_options, state="readonly", width=18)
        penalty_dropdown.grid(row=5, column=1, sticky="w", padx=5, pady=5)
        
        # penalty time input
        penalty_time_label = tk.Label(timer_frame, text="Time of Penalty:", bg="white")
        penalty_time_label.grid(row=6, column=0, sticky="w", padx=5, pady=5)
        penalty_time_entry = tk.Entry(timer_frame, width=20)
        penalty_time_entry.grid(row=6, column=1, sticky="w", padx=5, pady=5)
        
        # store references to widgets
        self.timer_frames[index] = {
            "frame": timer_frame,
            "player_entry": player_entry,
            "team_entry": team_entry,
            "time_var": time_var,
            "time_display": time_display,
            "penalty_var": penalty_var,
            "penalty_time_entry": penalty_time_entry,
            "close_btn": close_btn  # store reference to close button
        }
        
        # initialize timer values
        self.paused_times[index] = 0

def main():
    root = tk.Tk()
    app = LacrosseTimerApp(root)
    
    # start the app
    root.mainloop()

if __name__ == "__main__":
    main()
  
